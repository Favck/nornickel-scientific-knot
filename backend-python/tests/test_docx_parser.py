import os
import pytest
import docx
from parser.docx_parser import DocxParser


@pytest.fixture
def valid_docx_file(tmp_path):
    """
    Фикстура динамически создает DOCX-файл во временной директории.
    Файл содержит текст и таблицу со спецсимволами.
    """

    doc = docx.Document()

    doc.add_paragraph("Отчет по качеству сточных вод за 2024 год.")
    doc.add_paragraph(" ")

    table = doc.add_table(rows=2, cols=2)

    row0_cells = table.rows[0].cells
    row0_cells[0].text = "Параметр"
    row0_cells[1].text = "Норматив"

    row1_cells = table.rows[1].cells
    row1_cells[0].text = "сульфаты"
    row1_cells[1].text = "<=200 мг/л"

    file_path = tmp_path / "test_report.docx"
    doc.save(file_path)

    return str(file_path)


class TestDocxParser:
    """Тесты для парсера DOCX."""

    def test_parser_success_extraction(self, valid_docx_file):
        """Проверяет успешное извлечение текста, таблиц и установку статуса."""
        parser = DocxParser()
        result = parser.parse(valid_docx_file)

        # Проверяем метаданные
        assert (
            result["metadata"]["status"] == "success"
            ), "Статус должен быть 'success'"
        assert result["metadata"]["file_name"] == "test_report.docx"
        assert "error" not in result["metadata"]

        # Проверяем извлечение параграфов
        assert "Отчет по качеству сточных вод за 2024 год." in result["text"]

        # Проверяем склейку таблиц
        assert "Параметр | Норматив" in result["text"]
        assert (
            "сульфаты | <=200 мг/л" in result["text"]
        ), "Спецсимволы или единицы измерения потеряны!"

    def test_parser_file_not_found(self, tmp_path):
        """Проверяет, как парсер справляется с несуществующим файлом."""
        parser = DocxParser()
        fake_path = str(tmp_path / "does_not_exist.docx")

        result = parser.parse(fake_path)

        assert result["metadata"]["status"] == "error"
        assert (
            "error" in result["metadata"]
        ), "Должно присутствовать описание ошибки"
        assert result["text"] == "", "При ошибке текст должен быть пустым"
